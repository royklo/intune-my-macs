#!/usr/bin/env python3
# Copyright (c) Microsoft Corporation.
# Licensed under the MIT License.
"""
generate_payload_docs.py

Parses Settings Catalog / configuration policy JSON files and mobileconfig (.mobileconfig) plist files
to produce a markdown document describing the payloads and their key settings.

Output: INTUNE-MY-MACS-DOCUMENTATION.md at repo root.

Heuristics:
 - JSON Intune Settings Catalog: look for top-level keys like 'name', 'description', 'platforms', 'settings'.
 - 'settings' is an array of settingInstance containers. We traverse nested 'children' arrays.
 - Extract settingDefinitionId and any simple/choice values (value, string/integer) summarizing them.
 - Mobileconfig: parse plist, enumerate PayloadContent entries. For each payload dict list key/value excluding payload metadata keys.
 - Large payloads (> 60 entries) are truncated with summary note.
 - Provide both aggregated table and per-policy section with bullet list.

"""

from __future__ import annotations
import json
import plistlib
import pathlib
import re
import argparse
import xml.etree.ElementTree as ET
from typing import Any, Dict, List, Tuple

REPO_ROOT = pathlib.Path(__file__).resolve().parent.parent
OUTPUT_FILE = REPO_ROOT / "INTUNE-MY-MACS-DOCUMENTATION.md"
DOCX_OUTPUT_FILE = REPO_ROOT / "INTUNE-MY-MACS-DOCUMENTATION.docx"

JSON_GLOB = [
    "macOS/configurations/**/*.json",
    "macOS/mde/*.json",
]
MOBILECONFIG_GLOB = [
    "macOS/configurations/**/*.mobileconfig",
    "macOS/mde/*.mobileconfig",
]

METADATA_KEYS = {"PayloadDisplayName", "PayloadIdentifier", "PayloadType", "PayloadUUID", "PayloadVersion"}

def gather_files(patterns: List[str], suffix: str = "") -> List[pathlib.Path]:
    files: List[pathlib.Path] = []
    for pattern in patterns:
        for p in REPO_ROOT.glob(pattern):
            if p.is_file():
                if suffix and not p.name.endswith(suffix):
                    continue
                files.append(p)
    return sorted(files)

def safe_read_json(path: pathlib.Path) -> Dict[str, Any] | None:
    """Read JSON tolerating UTF-8 BOM."""
    try:
        # Read raw then decode handling BOM if present
        raw = path.read_bytes()
        text = raw.decode("utf-8-sig")  # utf-8-sig strips BOM if present
        return json.loads(text)
    except Exception as e:
        print(f"[WARN] Failed to parse JSON {path}: {e}")
        return None

def safe_read_plist(path: pathlib.Path) -> Dict[str, Any] | None:
    try:
        with path.open("rb") as f:
            return plistlib.load(f)
    except Exception as e:
        print(f"[WARN] Failed to parse mobileconfig plist {path}: {e}")
        return None

def extract_settings_catalog(json_doc: Dict[str, Any]) -> List[Tuple[str, str]]:
    """Return list of (settingDefinitionId, value) pairs by deep traversal.
    Handles nested settingInstance and groupSettingCollectionValue/children structures.
    """
    out: List[Tuple[str, str]] = []
    settings = json_doc.get("settings", [])

    def walk(node: Any):
        if isinstance(node, dict):
            sdid = node.get("settingDefinitionId")
            if sdid:
                # Extract choice value
                choice_block = node.get("choiceSettingValue")
                if isinstance(choice_block, dict):
                    choice_val = choice_block.get("value")
                    if choice_val is not None:
                        out.append((sdid, simplify_display_value(sdid, choice_val)))
                # Extract simple value
                simple_val_block = node.get("simpleSettingValue")
                if isinstance(simple_val_block, dict):
                    val = simple_val_block.get("value")
                    if val is not None:
                        out.append((sdid, simplify_display_value(sdid, val)))
                # Extract collection values (arrays)
                collection_block = node.get("simpleSettingCollectionValue")
                if isinstance(collection_block, list):
                    for idx, item in enumerate(collection_block):
                        if isinstance(item, dict):
                            val = item.get("value")
                            if val is not None:
                                # Use index suffix for multiple values
                                out.append((f"{sdid}[{idx}]", simplify_display_value(sdid, val)))
            # Recurse into all values to ensure we visit settingInstance etc.
            for v in node.values():
                walk(v)
        elif isinstance(node, list):
            for item in node:
                walk(item)

    walk(settings)
    return out

def simplify_value(val: Any) -> str:
    """Basic normalization converting numbers and truncating long strings."""
    if val is None:
        return ""
    if isinstance(val, (int, float)):
        return str(val)
    if isinstance(val, str):
        if len(val) > 120:
            return val[:117] + "..."
        return val
    return json.dumps(val)

BOOLEAN_SUFFIXES = {"_true": "True", "_false": "False"}

def simplify_display_value(key: str, raw_val: Any) -> str:
    """Remove duplicated key prefix embedded in Settings Catalog choice/simple values.
    Examples:
      key='com.apple.mcx.filevault2_enable', value='com.apple.mcx.filevault2_enable_0' -> '0'
      key='com.apple.systempolicy.control_EnableAssessment', value='com.apple.systempolicy.control_EnableAssessment_true' -> 'True'
      key='com.apple.managedclient.preferences_channelname', value='com.apple.managedclient.preferences_channelname_0' -> '0'
    Also converts *_true/*_false suffixes to True/False.
    Leaves placeholder tokens like {{mail}} intact.
    """
    # First apply base simplification
    val = simplify_value(raw_val)
    if not isinstance(val, str):
        return val
    # Preserve Jinja/placeholder tokens
    if re.search(r"{{.*?}}", val):
        return val
    lower_val = val.lower()
    # Map boolean suffix if present
    for suf, mapped in BOOLEAN_SUFFIXES.items():
        if lower_val.endswith(suf):
            # Only if preceding part matches key prefix to avoid accidental mapping
            prefix_part = val[: -len(suf)]
            if key.startswith(prefix_part.split('.')[-1]) or prefix_part.endswith(key.split('.')[-1]):
                return mapped
            # Common case: exact duplication of key then suffix
            if prefix_part == key:
                return mapped
    # If value starts with the key, strip that prefix plus underscore
    if val.startswith(key + "_"):
        trimmed = val[len(key) + 1:]
        return trimmed
    # Some values fully repeat path segments with underscores; attempt to drop matching leading portion
    if val.startswith(key):
        remainder = val[len(key):]
        remainder = remainder.lstrip('_')
        if remainder:
            return remainder
    return val

def extract_mobileconfig(plist_doc: Dict[str, Any]) -> List[Tuple[str, str]]:
    out: List[Tuple[str, str]] = []
    payloads = plist_doc.get("PayloadContent", [])
    for payload in payloads:
        if not isinstance(payload, dict):
            continue
        # Use PayloadType as prefix (more meaningful than DisplayName)
        prefix = payload.get("PayloadType", "payload")
        for k, v in payload.items():
            if k in METADATA_KEYS:
                continue
            if isinstance(v, (str, int, float)):
                out.append((f"{prefix}.{k}", simplify_value(v)))
            elif isinstance(v, (list, dict)):
                # Summarize complex structure size
                out.append((f"{prefix}.{k}", f"complex:{type(v).__name__}"))
    return out

def extract_compliance_policy(json_doc: Dict[str, Any]) -> List[Tuple[str, str]]:
    """Extract settings from compliance policy JSON (flat structure).
    Ignores metadata fields and extracts policy configuration.
    """
    out: List[Tuple[str, str]] = []
    IGNORE_KEYS = {
        "@odata.type", "displayName", "description", "id", "createdDateTime", 
        "lastModifiedDateTime", "version", "roleScopeTagIds"
    }
    
    for key, value in json_doc.items():
        if key in IGNORE_KEYS:
            continue
        # Handle scheduledActionsForRule separately
        if key == "scheduledActionsForRule" and isinstance(value, list):
            for idx, rule in enumerate(value):
                if isinstance(rule, dict):
                    rule_name = rule.get("ruleName", f"rule_{idx}")
                    configs = rule.get("scheduledActionConfigurations", [])
                    if configs:
                        out.append((f"{key}.{rule_name}.actionCount", str(len(configs))))
                        for cidx, config in enumerate(configs):
                            if isinstance(config, dict):
                                action_type = config.get("actionType", "unknown")
                                grace = config.get("gracePeriodHours", 0)
                                out.append((f"{key}.{rule_name}.action_{cidx}", f"{action_type} (grace: {grace}h)"))
        else:
            out.append((key, simplify_value(value)))
    
    return out

def format_table(rows: List[Tuple[str, str]]) -> str:
    """Return a markdown table with all rows (no truncation)."""
    if not rows:
        return "_No payload settings discovered_\n"
    header = "| Key | Value |\n|-----|-------|\n"
    body = "".join(f"| `{k}` | `{v}` |\n" for k, v in rows)
    return header + body

def classify_type(path: pathlib.Path) -> str:
    name = path.name
    # Derive from filename prefix (ref ID)
    m = re.match(r"([a-z]{3})-([a-z]{3})-(\d{3})", name)
    if m:
        prefix = m.group(1)
        mapping = {
            "pol": "Policy",
            "cfg": "CustomConfig",
            "cmp": "Compliance",
            "scr": "Script",
            "cat": "CustomAttribute",
            "app": "Package",
        }
        return mapping.get(prefix, "Unknown")
    # fallback
    if name.endswith(".mobileconfig"):
        return "CustomConfig"
    return "Policy"

def generate_markdown(entries: List[Dict[str, Any]]) -> str:
    md: List[str] = []
    import datetime
    today = datetime.date.today().strftime("%B %d, %Y")
    
    # Page 1: Cover Page (Large, Bold)
    md.append("# Intune My Macs\n\n")
    md.append("## Configuration Documentation\n\n")
    md.append(f"**Generated:** {today}\n\n")
    md.append(f"**Total Artifacts:** {len(entries)}\n\n")
    
    # Page 2: Project Description (Standard font)
    md.append("# About Intune My Macs\n\n")
    md.append("> **Proof of Concept \u2014 not for production use.** This repository is published as sample code to help teams evaluate and learn Microsoft Intune for macOS. The configurations and scripts are not a hardened baseline, are provided as-is without warranty or support, and must be reviewed, tested, and adapted before being deployed to managed devices.\n\n")
    md.append("**Intune My Macs** is a proof-of-concept configuration repository for Microsoft Intune-based macOS device management. ")
    md.append("This project provides sample policies, configuration profiles, scripts, and packages to help you evaluate and learn macOS device management with Intune. ")
    md.append("It is not a production baseline \u2014 review and adapt every artifact before deploying to managed devices.\n\n")
    md.append("## What's Included\n\n")
    md.append("This repository contains the following artifact types:\n\n")
    md.append("- **Settings Catalog Policies** - Modern declarative configuration policies\n")
    md.append("- **Custom Configuration Profiles** - Traditional mobileconfig profiles\n")
    md.append("- **Compliance Policies** - Device compliance requirements\n")
    md.append("- **Shell Scripts** - Automated configuration and remediation scripts\n")
    md.append("- **Application Packages** - macOS application installers\n")
    md.append("- **Custom Attributes** - Device inventory attributes\n\n")
    md.append("## About This Documentation\n\n")
    md.append("This document catalogs all configuration artifacts with complete settings details. ")
    md.append("Use the Index to quickly locate specific configurations, then refer to the detailed sections for complete settings breakdowns.\n\n")
    
    # Page 3: Index with Summary Table
    md.append("# Index\n\n")
    md.append("Click any reference ID to jump to detailed configuration.\n\n")
    md.append("| Ref | Type | Settings Count |\n|-----|------|----------------|\n")
    def anchor_for(ref: str, type_: str) -> str:
        # Mirror the heading line: ### ref (Type) -> pandoc/github anchor generation heuristic
        anchor_base = f"{ref}-{type_.lower()}"
        return anchor_base.replace(' ', '-').lower()
    for e in entries:
        anchor = anchor_for(e['ref'], e['type'])
        md.append(f"| [{e['ref']}](#{anchor}) | {e['type']} | {e['count']} |\n")
    md.append("\n")
    
    md.append("# Detailed Configuration\n\n")
    for e in entries:
        md.append(f"### {e['ref']} ({e['type']})\n\n")
        if e.get("description"):
            md.append(f"{e['description']}\n\n")
        md.append(f"**Source:** `{e['relpath']}`  \n")
        md.append(f"**Settings:** {e['count']}\n\n")
        md.append(format_table(e['settings']))
        md.append("\n\n")
    return "".join(md)

def add_page_breaks_for_docx(markdown: str) -> str:
    """Add OpenXML page breaks to markdown for Word/pandoc conversion.
    
    These page breaks are only used when generating DOCX files and should
    never appear in the committed markdown file.
    """
    PAGE_BREAK = "```{=openxml}\n<w:p><w:r><w:br w:type=\"page\"/></w:r></w:p>\n```\n\n"
    
    # Insert page break after "Total Artifacts" line (end of cover page)
    markdown = re.sub(
        r'(\*\*Total Artifacts:\*\* \d+\n\n)',
        r'\1' + PAGE_BREAK,
        markdown
    )
    
    # Insert page break after "About This Documentation" section (before Index)
    markdown = re.sub(
        r'(Use the Index to quickly locate specific configurations, then refer to the detailed sections for complete settings breakdowns\.\n\n)',
        r'\1' + PAGE_BREAK,
        markdown
    )
    
    # Insert page break after Index table (before Detailed Configuration)
    markdown = re.sub(
        r'(\| \[scr-sys-101-configure-dock\].*?\| Script \| \d+ \|\n\n)(# Detailed Configuration)',
        r'\1' + PAGE_BREAK + r'\2',
        markdown
    )
    
    return markdown

def markdown_to_docx(md_text: str, docx_path: pathlib.Path) -> None:
    """Very lightweight markdown to docx conversion focusing on headings, paragraphs, code spans and tables.
    Requires python-docx.
    """
    try:
        from docx import Document
    except ImportError:
        print("[WARN] python-docx not installed; skipping DOCX generation")
        return

    document = Document()
    lines = md_text.splitlines()
    table_buffer: List[List[str]] = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        headers = []
        if len(table_buffer) >= 2 and set(table_buffer[1][0]) == {'-'}:
            headers = table_buffer[0]
            data_rows = table_buffer[2:]
        else:
            headers = table_buffer[0]
            data_rows = table_buffer[1:]
        tbl = document.add_table(rows=1, cols=len(headers))
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass  # style may not exist in some environments
        # Header formatting
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h.strip('| ').strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        # Apply shading to header row (light gray)
        hdr_tr = tbl.rows[0]._tr
        for tc in hdr_tr.tc_lst:
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')  # light gray
            tcPr.append(shd)
        # Data rows
        for r in data_rows:
            row = tbl.add_row()
            for i, c in enumerate(r):
                cell = row.cells[i]
                text = c.strip('| ').strip()
                # Remove backticks for cleaner display
                text = text.replace('`', '')
                cell.text = text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        # Monospace for first column (Key)
                        if i == 0:
                            run.font.name = 'Courier New'
        # Auto-fit attempt (not always respected)
        for col in tbl.columns:
            for cell in col.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.size is None:
                            run.font.size = Pt(9)
        table_buffer = []

    for line in lines:
        if re.match(r'^\s*\|', line):
            # Part of a table
            table_buffer.append([c for c in line.split('|') if c])
            continue
        else:
            flush_table()
        if line.startswith('# '):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith('---'):
            document.add_page_break()
        elif line.strip():
            # Inline code spans: replace `text` with just text (could style later)
            cleaned = re.sub(r'`([^`]+)`', r'\1', line)
            document.add_paragraph(cleaned)
    flush_table()
    document.save(str(docx_path))
    print(f"[INFO] Wrote DOCX to {docx_path}")

def build_entries(include_mde: bool = False) -> List[Dict[str, Any]]:
    json_files = gather_files(JSON_GLOB, suffix=".json")
    mc_files = gather_files(MOBILECONFIG_GLOB, suffix=".mobileconfig")
    
    # Filter out MDE folder unless --mde flag is passed
    if not include_mde:
        json_files = [f for f in json_files if not str(f).startswith(str(REPO_ROOT / "macOS" / "mde"))]
        mc_files = [f for f in mc_files if not str(f).startswith(str(REPO_ROOT / "macOS" / "mde"))]
    
    entries: List[Dict[str, Any]] = []
    # Helper: try to load manifest XML next to source file (same base name)
    def load_manifest_metadata(source_path: pathlib.Path) -> Dict[str, str]:
        meta: Dict[str, str] = {}
        manifest_path = source_path.with_suffix('.xml')
        if manifest_path.exists():
            try:
                tree = ET.parse(manifest_path)
                root = tree.getroot()
                name_el = root.find('Name')
                desc_el = root.find('Description')
                type_el = root.find('Type')
                if name_el is not None and name_el.text:
                    meta['name'] = name_el.text.strip()
                if desc_el is not None and desc_el.text:
                    meta['description'] = desc_el.text.strip()
                if type_el is not None and type_el.text:
                    meta['type'] = type_el.text.strip()
            except Exception as e:
                print(f"[WARN] Failed to parse manifest XML {manifest_path}: {e}")
        return meta
    for f in json_files:
        doc = safe_read_json(f)
        if not doc:
            continue
        
        # Determine policy type and extract settings accordingly
        odata_type = doc.get("@odata.type", "")
        settings = []
        
        # Try Settings Catalog format first
        settings = extract_settings_catalog(doc)
        
        # If no settings found, check for compliance policy
        if not settings and "CompliancePolicy" in odata_type:
            settings = extract_compliance_policy(doc)
        
        # Fallback for enrollment restriction style JSON
        if not settings and odata_type.endswith("deviceEnrollmentPlatformRestriction"):
            pr = doc.get("platformRestriction", {})
            if isinstance(pr, dict):
                for k, v in pr.items():
                    settings.append((f"platformRestriction.{k}", simplify_value(v)))
        
        rel = f.relative_to(REPO_ROOT)
        ref_id = f.stem
        manifest_meta = load_manifest_metadata(f)
        derived_type = classify_type(f)
        if 'type' in manifest_meta:
            derived_type = manifest_meta['type']
        entries.append({
            "ref": ref_id,
            "type": derived_type,
            "relpath": str(rel),
            "name": manifest_meta.get("name"),
            "description": manifest_meta.get("description"),
            "settings": settings,
            "count": len(settings)
        })

    # Add standalone manifests for Package, Script, CustomAttribute not covered above
    # We discover all XML manifests and include those whose SourceFile points to a .pkg/.sh/.zsh etc.
    xml_manifests = list(REPO_ROOT.rglob("*.xml"))
    
    # Filter out MDE folder unless include_mde is True
    if not include_mde:
        xml_manifests = [m for m in xml_manifests if not str(m).startswith(str(REPO_ROOT / "macOS" / "mde"))]
    
    for mpath in xml_manifests:
        try:
            tree = ET.parse(mpath)
            root = tree.getroot()
            if root.tag not in ('MacIntuneManifest', 'IntuneManifest'):
                continue
            type_el = root.find('Type')
            src_el = root.find('SourceFile')
            name_el = root.find('Name')
            desc_el = root.find('Description')
            ref_id_el = root.find('ReferenceId')
            if type_el is None or src_el is None:
                continue
            artifact_type = type_el.text.strip()
            rel_source = src_el.text.strip()
            
            # Skip if already processed:
            # - Policy/CustomConfig/Compliance that point to .json files (handled by JSON processing)
            # - CustomConfig that points to .mobileconfig (handled by plist processing)
            # These should ALWAYS be skipped since JSON/mobileconfig processing happens first
            if artifact_type in {'Policy', 'CustomConfig', 'Compliance'} and rel_source.endswith('.json'):
                continue
            if artifact_type == 'CustomConfig' and rel_source.endswith('.mobileconfig'):
                continue
            
            # SourceFile is relative to the manifest's platform folder (e.g. macOS/)
            platform_root = mpath.relative_to(REPO_ROOT).parts[0]
            rel_source = f"{platform_root}/{rel_source}"
            # Additional check: skip if already in entries by relpath
            rel_path_obj = REPO_ROOT / rel_source
            already = any(e['relpath'] == rel_source for e in entries)
            if already:
                continue
            # Extract subtree settings for Script, Package, CustomAttribute
            settings: List[Tuple[str, str]] = []
            if artifact_type in {"Script", "Package", "CustomAttribute"}:
                subtree_tag = {
                    "Script": "Script",
                    "Package": "Package",
                    "CustomAttribute": "CustomAttribute"
                }[artifact_type]
                subtree = root.find(subtree_tag)
                if subtree is not None:
                    for child in list(subtree):
                        if child.text:
                            settings.append((child.tag, child.text.strip()))
            entries.append({
                "ref": (rel_path_obj.stem if rel_path_obj.exists() else (mpath.stem)),
                "type": artifact_type,
                "relpath": rel_source,
                "name": name_el.text.strip() if name_el is not None and name_el.text else None,
                "description": desc_el.text.strip() if desc_el is not None and desc_el.text else "",
                "settings": settings,
                "count": len(settings),
            })
        except Exception as e:
            print(f"[WARN] Failed processing manifest {mpath}: {e}")
    for f in mc_files:
        doc = safe_read_plist(f)
        if not doc:
            continue
        settings = extract_mobileconfig(doc)
        rel = f.relative_to(REPO_ROOT)
        ref_id = f.stem
        manifest_meta = load_manifest_metadata(f)
        derived_type = classify_type(f)
        if 'type' in manifest_meta:
            derived_type = manifest_meta['type']
        entries.append({
            "ref": ref_id,
            "type": derived_type,
            "relpath": str(rel),
            "name": manifest_meta.get("name") or doc.get("PayloadDisplayName"),
            "description": manifest_meta.get("description", ""),
            "settings": settings,
            "count": len(settings)
        })
    
    # Deduplicate entries by (ref, type, relpath) tuple
    seen = set()
    deduped = []
    for entry in entries:
        key = (entry['ref'], entry['type'], entry['relpath'])
        if key not in seen:
            seen.add(key)
            deduped.append(entry)
    
    deduped.sort(key=lambda x: x['ref'])
    return deduped

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate payload documentation (Markdown + optional DOCX)")
    parser.add_argument("--docx", action="store_true", help="Also generate a DOCX file")
    parser.add_argument("--pandoc", action="store_true", help="Use pandoc for DOCX conversion (requires pandoc installed)")
    parser.add_argument("--mde", action="store_true", help="Include MDE (Microsoft Defender for Endpoint) folder in documentation")
    args = parser.parse_args()

    entries = build_entries(include_mde=args.mde)
    markdown = generate_markdown(entries)
    OUTPUT_FILE.write_text(markdown, encoding="utf-8")
    print(f"[INFO] Wrote markdown to {OUTPUT_FILE}")
    print(f"[INFO] Documented {len(entries)} payload artifacts")
    if args.docx:
        if args.pandoc:
            # Attempt pandoc conversion
            import shutil, subprocess, tempfile
            pandoc_exe = shutil.which("pandoc")
            if not pandoc_exe:
                print("[WARN] --pandoc requested but pandoc not found; falling back to internal converter")
                markdown_to_docx(markdown, DOCX_OUTPUT_FILE)
            else:
                try:
                    # Create markdown with page breaks for Word conversion
                    markdown_with_breaks = add_page_breaks_for_docx(markdown)
                    # Write temp markdown file, run pandoc
                    with tempfile.NamedTemporaryFile(suffix='.md', delete=False) as tmp_md:
                        tmp_md.write(markdown_with_breaks.encode('utf-8'))
                        tmp_md_path = tmp_md.name
                    cmd = [pandoc_exe, '-f', 'markdown', tmp_md_path, '-o', str(DOCX_OUTPUT_FILE), '--standalone']
                    print(f"[INFO] Running pandoc: {' '.join(cmd)}")
                    subprocess.run(cmd, check=True)
                    print(f"[INFO] Wrote DOCX via pandoc to {DOCX_OUTPUT_FILE}")
                    # Post-process tables for styling
                    try:
                        from docx import Document
                        from docx.shared import Pt
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        doc = Document(str(DOCX_OUTPUT_FILE))
                        table_count = 0
                        for tbl_idx, tbl in enumerate(doc.tables):
                            table_count += 1
                            # Apply grid style if exists
                            try:
                                tbl.style = 'Table Grid'
                            except Exception:
                                pass
                            # First table is summary table - use smaller font (9pt) for all cells
                            is_summary_table = (tbl_idx == 0)
                            # Set table to autofit contents
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            tbl_element = tbl._tbl
                            tblPr = tbl_element.tblPr
                            if tblPr is None:
                                tblPr = OxmlElement('w:tblPr')
                                tbl_element.insert(0, tblPr)
                            
                            # Remove any existing table width setting
                            for e in list(tblPr):
                                if e.tag == qn('w:tblW'):
                                    tblPr.remove(e)
                            
                            # Set table width to AUTO (0) to enable autofit to contents
                            tblW = OxmlElement('w:tblW')
                            tblW.set(qn('w:w'), '0')
                            tblW.set(qn('w:type'), 'auto')
                            tblPr.append(tblW)
                            
                            # Set autofit layout
                            for e in list(tblPr):
                                if e.tag == qn('w:tblLayout'):
                                    tblPr.remove(e)
                            layout = OxmlElement('w:tblLayout')
                            layout.set(qn('w:type'), 'autofit')
                            tblPr.append(layout)
                            
                            # Remove fixed width constraints on all columns to allow autofit
                            tblGrid = tbl_element.find(qn('w:tblGrid'))
                            if tblGrid is not None:
                                for gridCol in tblGrid.findall(qn('w:gridCol')):
                                    # Remove w:w attribute (fixed width)
                                    if qn('w:w') in gridCol.attrib:
                                        del gridCol.attrib[qn('w:w')]
                            
                            # Remove cell width constraints
                            for row in tbl._element.findall(qn('w:tr')):
                                for tc in row.findall(qn('w:tc')):
                                    tcPr = tc.find(qn('w:tcPr'))
                                    if tcPr is not None:
                                        tcW = tcPr.find(qn('w:tcW'))
                                        if tcW is not None:
                                            tcPr.remove(tcW)
                            # Enforce borders (grid) even if style not applied
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            # Access or create tblPr for borders
                            tbl_element = tbl._tbl
                            tblPr = tbl_element.tblPr
                            if tblPr is None:
                                tblPr = OxmlElement('w:tblPr')
                                tbl_element.insert(0, tblPr)
                            # Remove existing borders then set new
                            for e in list(tblPr):
                                if e.tag == qn('w:tblBorders'):
                                    tblPr.remove(e)
                            borders = OxmlElement('w:tblBorders')
                            for side in ['top','left','bottom','right','insideH','insideV']:
                                elem = OxmlElement(f'w:{side}')
                                elem.set(qn('w:val'), 'single')
                                elem.set(qn('w:sz'), '6')  # ~0.5pt
                                elem.set(qn('w:space'), '0')
                                elem.set(qn('w:color'), '000000')
                                borders.append(elem)
                            tblPr.append(borders)
                            if tbl.rows:
                                # Header row shading & bold
                                hdr = tbl.rows[0]
                                for ci, cell in enumerate(hdr.cells):
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.font.bold = True
                                            run.font.name = 'Courier New'
                                            run.font.size = Pt(8)
                                    tc = cell._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'), 'D9D9D9')
                                    tcPr.append(shd)
                            # Data rows formatting
                            for r_index, row in enumerate(tbl.rows[1:], start=1):
                                for ci, cell in enumerate(row.cells):
                                    for p in cell.paragraphs:
                                        # Also check for hyperlinks which have separate styling
                                        for hyperlink in p._element.findall('.//' + qn('w:hyperlink')):
                                            for run_elem in hyperlink.findall('.//' + qn('w:r')):
                                                rPr = run_elem.find(qn('w:rPr'))
                                                if rPr is None:
                                                    rPr = OxmlElement('w:rPr')
                                                    run_elem.insert(0, rPr)
                                                # Remove existing font size
                                                for sz in list(rPr.findall(qn('w:sz'))):
                                                    rPr.remove(sz)
                                                for szCs in list(rPr.findall(qn('w:szCs'))):
                                                    rPr.remove(szCs)
                                                # Set new font size
                                                if is_summary_table:
                                                    sz = OxmlElement('w:sz')
                                                    sz.set(qn('w:val'), '16')  # 8pt = 16 half-points
                                                    rPr.append(sz)
                                                    szCs = OxmlElement('w:szCs')
                                                    szCs.set(qn('w:val'), '16')
                                                    rPr.append(szCs)
                                        
                                        for run in p.runs:
                                            # All table text uses Courier New 8pt
                                            run.font.name = 'Courier New'
                                            run.font.size = Pt(8)
                                            # Value column (second column, ci==1) should be bold
                                            if ci == 1:
                                                run.font.bold = True
                        
                        # Set paragraph and heading fonts to Aptos
                        # Find first 3 headings (cover page sections) and make them large/bold
                        h1_count = 0
                        for para in doc.paragraphs:
                            # Count H1 headings to track pages (0=cover, 1=description, 2=index, 3+=details)
                            if para.style.name == 'Heading 1':
                                h1_count += 1
                            
                            is_cover_page = (h1_count == 1)  # First H1 is cover page
                            
                            for run in para.runs:
                                run.font.name = 'Aptos'
                                run.font.size = Pt(11)
                                # Make cover page text bold
                                if is_cover_page:
                                    run.font.bold = True
                            
                            # Headings get larger font
                            if para.style.name.startswith('Heading'):
                                para.paragraph_format.space_after = Pt(8)
                                para.paragraph_format.space_before = Pt(12)
                                for run in para.runs:
                                    if para.style.name == 'Heading 1':
                                        # Cover page gets huge fonts
                                        run.font.size = Pt(36 if is_cover_page else 16)
                                    elif para.style.name == 'Heading 2':
                                        run.font.size = Pt(24 if is_cover_page else 14)
                                    else:
                                        run.font.size = Pt(14)
                        
                        # Set document to modern Word format (removes compatibility mode)
                        # Create new compat settings for Word 2016+
                        settings_element = doc.settings.element
                        
                        # Remove old compatibility settings
                        for compat in list(settings_element.findall(qn('w:compat'))):
                            settings_element.remove(compat)
                        
                        # Add modern compatibility mode settings
                        compat = OxmlElement('w:compat')
                        # Set compatibilityMode to 15 (Word 2013+) or 16 (Word 2016+)
                        compat_setting = OxmlElement('w:compatSetting')
                        compat_setting.set(qn('w:name'), 'compatibilityMode')
                        compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
                        compat_setting.set(qn('w:val'), '16')  # Word 2016+ format
                        compat.append(compat_setting)
                        settings_element.append(compat)
                        
                        doc.save(str(DOCX_OUTPUT_FILE))
                        print(f"[INFO] Post-processed {table_count} tables in pandoc DOCX (autofit=on by default)")
                    except Exception as e:
                        print(f"[WARN] DOCX post-processing failed: {e}")
                except subprocess.CalledProcessError as e:
                    print(f"[WARN] pandoc failed ({e}); falling back to internal converter")
                    markdown_to_docx(markdown, DOCX_OUTPUT_FILE)
                finally:
                    try:
                        pathlib.Path(tmp_md_path).unlink(missing_ok=True)
                    except Exception:
                        pass
        else:
            markdown_to_docx(markdown, DOCX_OUTPUT_FILE)

if __name__ == "__main__":
    main()
